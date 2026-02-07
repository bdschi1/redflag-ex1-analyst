"""
document_loader.py

Unified document loading for the RedFlag Analyst engine.
Accepts .txt, .pdf, and .docx files and extracts plain text.

Design goals:
- Keep redflag_engine.py accepting `text: str` — this module sits upstream.
- Local-first, no API keys.
- Backward-compatible: TXT loading produces identical output to the legacy path.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class LoadResult:
    """Result of loading a document."""

    text: str
    source_path: str | None  # None when loaded from bytes
    format: str  # "txt", "pdf", "docx"
    page_count: int | None  # Meaningful for PDF only
    char_count: int = 0
    warnings: list[str] = field(default_factory=list)

    def __post_init__(self) -> None:
        self.char_count = len(self.text)


class UnsupportedFormatError(Exception):
    """Raised when a document format is not supported."""


class DocumentLoader:
    """
    Load text content from supported document formats.

    Supported: .txt, .pdf, .docx
    Unsupported .doc files raise UnsupportedFormatError with conversion guidance.
    """

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}

    def load_file(self, path: str | Path) -> LoadResult:
        """Load text from a file on disk, dispatching by extension."""
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        ext = path.suffix.lower()
        self._validate_extension(ext)

        raw = path.read_bytes()
        result = self._dispatch(raw, ext)
        result.source_path = str(path)
        return result

    def load_bytes(self, data: bytes, filename: str) -> LoadResult:
        """Load text from in-memory bytes (e.g. Streamlit uploads)."""
        ext = Path(filename).suffix.lower()
        self._validate_extension(ext)

        result = self._dispatch(data, ext)
        result.source_path = None
        return result

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _validate_extension(self, ext: str) -> None:
        if ext == ".doc":
            raise UnsupportedFormatError(
                "Legacy .doc format is not supported. "
                "Please convert to .docx (Save As in Word/LibreOffice) or .pdf."
            )
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(
                f"Unsupported file format: '{ext}'. "
                f"Supported formats: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

    def _dispatch(self, data: bytes, ext: str) -> LoadResult:
        if ext == ".txt":
            return self._load_txt(data)
        if ext == ".pdf":
            return self._load_pdf(data)
        if ext == ".docx":
            return self._load_docx(data)
        # Should not reach here due to _validate_extension
        raise UnsupportedFormatError(f"Unhandled extension: {ext}")  # pragma: no cover

    def _load_txt(self, data: bytes) -> LoadResult:
        text = data.decode("utf-8", errors="replace")
        return LoadResult(
            text=text,
            source_path=None,
            format="txt",
            page_count=None,
        )

    def _load_pdf(self, data: bytes) -> LoadResult:
        import pdfplumber

        warnings: list[str] = []

        try:
            pdf = pdfplumber.open(io.BytesIO(data))
        except Exception as exc:
            raise ValueError(f"Failed to open PDF: {exc}") from exc

        pages_text: list[str] = []
        with pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)

        text = "\n\n".join(pages_text).strip()

        if not text:
            warnings.append(
                "PDF appears to contain no extractable text "
                "(scanned/image-only PDFs are not supported without OCR)."
            )

        return LoadResult(
            text=text,
            source_path=None,
            format="pdf",
            page_count=page_count,
            warnings=warnings,
        )

    def _load_docx(self, data: bytes) -> LoadResult:
        from docx import Document

        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:
            raise ValueError(f"Failed to open DOCX: {exc}") from exc

        parts: list[str] = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract table cell text
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts).strip()

        return LoadResult(
            text=text,
            source_path=None,
            format="docx",
            page_count=None,
        )
