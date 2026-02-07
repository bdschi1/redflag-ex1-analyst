"""
Shared test fixtures for RedFlag Analyst tests.

Generates PDF and DOCX test files dynamically (since .gitignore blocks *.pdf).
"""

from __future__ import annotations

import os
import tempfile

import pytest


@pytest.fixture
def tmp_dir():
    """Provide a temporary directory that is cleaned up after the test."""
    with tempfile.TemporaryDirectory() as d:
        yield d


@pytest.fixture
def sample_txt_path(tmp_dir):
    """Create a simple .txt file."""
    path = os.path.join(tmp_dir, "sample.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write("This is a clean analyst note about ACME Corp revenue growth.\n")
    return path


@pytest.fixture
def risky_txt_path(tmp_dir):
    """Create a .txt file with MNPI triggers."""
    path = os.path.join(tmp_dir, "risky.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(
            "After 15 calls with the expert, a friend (an investigator) "
            "said things look good for the trial.\n"
        )
    return path


@pytest.fixture
def sample_pdf_path(tmp_dir):
    """Generate a simple single-page PDF."""
    from fpdf import FPDF

    path = os.path.join(tmp_dir, "sample.pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(200, 10, text="Investment Thesis for ACME Corp", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(
        200,
        10,
        text="Revenue growth is strong based on public filings.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.output(path)
    return path


@pytest.fixture
def multi_page_pdf_path(tmp_dir):
    """Generate a multi-page PDF."""
    from fpdf import FPDF

    path = os.path.join(tmp_dir, "multi_page.pdf")
    pdf = FPDF()
    for i in range(3):
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(
            200,
            10,
            text=f"Page {i + 1}: Analysis section {i + 1}",
            new_x="LMARGIN",
            new_y="NEXT",
        )
        pdf.cell(
            200,
            10,
            text="Detailed analysis content here.",
            new_x="LMARGIN",
            new_y="NEXT",
        )
    pdf.output(path)
    return path


@pytest.fixture
def risky_pdf_path(tmp_dir):
    """Generate a PDF with MNPI risk content."""
    from fpdf import FPDF

    path = os.path.join(tmp_dir, "risky.pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(
        200,
        10,
        text="A friend who is an investigator said things look good.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.cell(
        200,
        10,
        text="Off the record, preliminary results are positive.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.output(path)
    return path


@pytest.fixture
def pdf_with_boilerplate_path(tmp_dir):
    """Generate a PDF with both risky content and boilerplate disclaimers."""
    from fpdf import FPDF

    path = os.path.join(tmp_dir, "boilerplate.pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    # Substantive risky content
    pdf.cell(
        200,
        10,
        text="Off the record, the insider confirmed the deal.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.cell(200, 10, text="", new_x="LMARGIN", new_y="NEXT")
    # Boilerplate
    pdf.cell(
        200,
        10,
        text="This report is for institutional investors only.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.cell(
        200,
        10,
        text="Past performance is not indicative of future results.",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.output(path)
    return path


@pytest.fixture
def sample_docx_path(tmp_dir):
    """Generate a simple .docx file."""
    from docx import Document

    path = os.path.join(tmp_dir, "sample.docx")
    doc = Document()
    doc.add_paragraph("Investment Thesis for ACME Corp")
    doc.add_paragraph("Revenue growth is strong based on public filings.")
    doc.save(path)
    return path


@pytest.fixture
def risky_docx_path(tmp_dir):
    """Generate a .docx with MNPI content."""
    from docx import Document

    path = os.path.join(tmp_dir, "risky.docx")
    doc = Document()
    doc.add_paragraph("A friend who is an investigator said things look good for the trial.")
    doc.add_paragraph("Off the record, preliminary results are positive.")
    doc.save(path)
    return path


@pytest.fixture
def docx_with_table_path(tmp_dir):
    """Generate a .docx with paragraphs and a table."""
    from docx import Document

    path = os.path.join(tmp_dir, "with_table.docx")
    doc = Document()
    doc.add_paragraph("Financial Summary")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "$1.5B"
    doc.save(path)
    return path


@pytest.fixture
def empty_txt_path(tmp_dir):
    """Create an empty .txt file."""
    path = os.path.join(tmp_dir, "empty.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write("")
    return path
